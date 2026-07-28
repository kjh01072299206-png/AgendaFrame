from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = ROOT / "outputs"
OUT = DEFAULT_OUT
DOCS = ROOT / "docs"
PLANNING = DOCS / "planning"
SPECS = DOCS / "specs"
RESEARCH = DOCS / "research"
SUBMISSION = DOCS / "submission"
DEFAULT_DOCX_OUT = SUBMISSION / "final-report.docx"


WBS_ROWS = [
    ["1.0", "기획 및 수집 기준", "MVP 범위와 PBL-01~20 백로그 확정", "공동", "수정 프로덕트 백로그", "2026-07-07", "2026-07-08", "2일", "-"],
    ["1.1", "기획 및 수집 기준", "분석 대상 언론사 3~5개 선정", "공동", "대상 언론사 목록", "2026-07-07", "2026-07-09", "3일", "1.0"],
    ["1.2", "기획 및 수집 기준", "정책 분야 기준과 수집 항목 정의", "최지우", "정책 분야·수집 항목 명세", "2026-07-08", "2026-07-11", "4일", "1.0"],
    ["1.3", "기획 및 수집 기준", "기사 전문 비저장·원문 링크 중심 저장 정책 정리", "최지우", "저작권 고려 저장 정책", "2026-07-09", "2026-07-12", "4일", "1.2"],
    ["2.0", "기사 수집 MVP", "언론사별 홈페이지 구조와 주요 섹션 조사", "강준혁", "언론사별 구조 분석표", "2026-07-10", "2026-07-14", "5일", "1.1"],
    ["2.1", "기사 수집 MVP", "Playwright 기반 기사 메타데이터 수집 MVP 구현", "강준혁", "수집 코드·샘플 데이터", "2026-07-14", "2026-07-23", "10일", "2.0"],
    ["2.2", "기사 수집 MVP", "3개 언론사 확장, 선택자 테스트, 수집 로그 구현", "강준혁", "수집 로그·갱신 시각 화면", "2026-07-21", "2026-07-31", "11일", "2.1"],
    ["3.0", "저장·이슈 묶기", "기사·언론사·수집 로그 DB 설계 및 저장 로직 구현", "강준혁", "메타데이터 DB 스키마", "2026-07-17", "2026-07-26", "10일", "1.3, 2.1"],
    ["3.1", "저장·이슈 묶기", "임베딩 기준 설계와 유사 기사 클러스터링 MVP 구현", "강준혁", "이슈 클러스터링 모듈", "2026-07-24", "2026-08-05", "13일", "3.0"],
    ["3.2", "저장·이슈 묶기", "클러스터링 검증 샘플 작성 및 정합성 비교", "공동", "클러스터링 검증표", "2026-08-01", "2026-08-08", "8일", "3.1"],
    ["4.0", "의제 점수·대시보드", "의제 점수 산식과 배치 위치 가중치 정의", "최지우", "의제 점수 산식표", "2026-08-01", "2026-08-08", "8일", "3.1"],
    ["4.1", "의제 점수·대시보드", "의제 점수 계산 로직과 랭킹 API/UI 구현", "강준혁", "오늘의 의제 랭킹 화면", "2026-08-05", "2026-08-14", "10일", "4.0"],
    ["4.2", "의제 점수·대시보드", "이슈 상세, 기사 목록, 원문 링크 이동 구현", "강준혁", "이슈 상세 화면", "2026-08-10", "2026-08-17", "8일", "3.1, 4.1"],
    ["4.3", "의제 점수·대시보드", "언론사별 보도 건수와 배치 위치 비교 구현", "강준혁", "언론사 비교 화면", "2026-08-14", "2026-08-21", "8일", "4.2"],
    ["5.0", "프레임·AI 리포트", "프레임 코드북과 근거 표시 정책 확정", "최지우", "프레임 코드북·근거 정책", "2026-08-05", "2026-08-14", "10일", "1.3"],
    ["5.1", "프레임·AI 리포트", "Gemini 프레임 분석 프롬프트/API/저장 구조 구현", "공동", "프레임 분석 결과", "2026-08-12", "2026-08-22", "11일", "5.0, 4.2"],
    ["5.2", "프레임·AI 리포트", "프레임 비중 그래프와 근거 표현 UI 구현", "강준혁", "프레임 비교 화면", "2026-08-18", "2026-08-25", "8일", "5.1"],
    ["5.3", "프레임·AI 리포트", "프레임 수동 검증 및 오류 유형 정리", "공동", "프레임 검증표", "2026-08-22", "2026-08-28", "7일", "5.1"],
    ["5.4", "프레임·AI 리포트", "AI 리포트 프롬프트/API/UI 구현", "공동", "AI 리포트 화면", "2026-08-20", "2026-08-29", "10일", "4.1, 5.1"],
    ["6.0", "검증 및 발표 준비", "수집부터 리포트까지 MVP 최종 시나리오 테스트", "공동", "통합 테스트 결과", "2026-08-26", "2026-08-31", "6일", "4.1~5.4"],
    ["6.1", "검증 및 발표 준비", "발표 자료와 데모 시나리오 작성", "공동", "발표 자료·데모 시나리오", "2026-09-01", "2026-10-08", "38일", "6.0"],
]


def split_revision_backlog(text: str) -> tuple[str, str, str]:
    product_start = text.index("## AgendaFrame 프로덕트 백로그")
    sprint_start = text.index("## AgendaFrame 스프린트 백로그")
    later_start = text.index("## 후순위 백로그")
    product = text[product_start:sprint_start].strip()
    sprint = text[sprint_start:later_start].strip()
    later = text[later_start:].strip()
    product = product.replace("## AgendaFrame 프로덕트 백로그", "# AgendaFrame 프로덕트 백로그", 1)
    sprint = sprint.replace("## AgendaFrame 스프린트 백로그", "# AgendaFrame 스프린트 백로그", 1)
    return product + "\n\n" + later + "\n", sprint + "\n\n" + later + "\n", later


def build_wbs_markdown() -> str:
    lines = [
        "# 10. AgendaFrame WBS 및 간트차트",
        "",
        "작성일: 2026-07-07  ",
        "작성 담당: 강준혁  ",
        "프로젝트명: AgendaFrame",
        "",
        "## 1. 프로젝트 개요",
        "",
        "AgendaFrame은 주요 언론사의 홈페이지 배치와 보도 빈도를 기반으로 오늘의 공적 의제를 산출하고, 동일 이슈에 대한 언론사별 관점/프레임 차이를 비교하는 AI 기반 뉴스 의제·프레임 분석 플랫폼이다.",
        "",
        "이번 WBS는 수정 백로그의 PBL-01~PBL-20 기준에 맞춰 MVP 핵심 범위와 후순위 확장 범위를 분리했다. MVP는 기사 메타데이터 수집, 기사 전문 비저장 저장 정책, 유사 이슈 클러스터링, 의제 점수 산출, 이슈 상세·언론사 비교, 프레임 분석, AI 리포트, 검증까지를 포함한다.",
        "",
        "## 2. WBS",
        "",
        "| WBS ID | 상위 작업 | 세부 작업 | 담당자 | 산출물 | 시작일 | 종료일 | 기간 | 선행 작업 |",
        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
    ]
    lines += ["| " + " | ".join(row) + " |" for row in WBS_ROWS]
    lines += [
        "",
        "## 3. 간트차트 표",
        "",
        "| 작업 | 7/7~7/13 | 7/14~7/20 | 7/21~7/27 | 7/28~8/3 | 8/4~8/10 | 8/11~8/17 | 8/18~8/24 | 8/25~8/31 | 9/1~10/8 |",
        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
        "| 범위·수집 기준 확정 | X |  |  |  |  |  |  |  |  |",
        "| 기사 수집 MVP |  | X | X |  |  |  |  |  |  |",
        "| 저장 정책·DB 설계 |  | X | X | X |  |  |  |  |  |",
        "| 이슈 클러스터링·초기 검증 |  |  | X | X | X |  |  |  |  |",
        "| 의제 점수·랭킹 대시보드 |  |  |  | X | X | X |  |  |  |",
        "| 이슈 상세·언론사 비교 |  |  |  |  | X | X | X |  |  |",
        "| 프레임 분석·근거 표시 |  |  |  |  | X | X | X | X |  |",
        "| AI 리포트 구현 |  |  |  |  |  | X | X | X |  |",
        "| MVP 통합 테스트 |  |  |  |  |  |  |  | X |  |",
        "| 발표 자료·데모 시나리오 |  |  |  |  |  |  |  |  | X |",
        "",
        "## 4. Mermaid 간트차트",
        "",
        "```mermaid",
        "gantt",
        "    title AgendaFrame Revised WBS Gantt Chart",
        "    dateFormat YYYY-MM-DD",
        "    axisFormat %m/%d",
        "",
        "    section 기획 및 수집 기준",
        "    MVP 범위와 백로그 확정             :a1, 2026-07-07, 2d",
        "    대상 언론사 선정                   :a2, 2026-07-07, 3d",
        "    수집 항목·저장 정책 정의           :a3, 2026-07-08, 5d",
        "",
        "    section 기사 수집 MVP",
        "    홈페이지 구조 조사                 :b1, 2026-07-10, 5d",
        "    Playwright 수집 MVP                :b2, 2026-07-14, 10d",
        "    3개 언론사 확장·수집 로그          :b3, 2026-07-21, 11d",
        "",
        "    section 저장·이슈 묶기",
        "    DB 설계·저장 로직                  :c1, 2026-07-17, 10d",
        "    임베딩·클러스터링 구현             :c2, 2026-07-24, 13d",
        "    클러스터링 검증                    :c3, 2026-08-01, 8d",
        "",
        "    section 의제 점수·대시보드",
        "    의제 점수 산식 정의                :d1, 2026-08-01, 8d",
        "    랭킹 API/UI 구현                   :d2, 2026-08-05, 10d",
        "    이슈 상세·언론사 비교              :d3, 2026-08-10, 12d",
        "",
        "    section 프레임·AI 리포트",
        "    프레임 코드북·근거 정책            :e1, 2026-08-05, 10d",
        "    Gemini 프레임 분석 구현            :e2, 2026-08-12, 11d",
        "    프레임 비교·수동 검증              :e3, 2026-08-18, 11d",
        "    AI 리포트 구현                     :e4, 2026-08-20, 10d",
        "",
        "    section 검증 및 발표 준비",
        "    MVP 통합 테스트                    :f1, 2026-08-26, 6d",
        "    발표 자료·데모 시나리오            :f2, 2026-09-01, 38d",
        "```",
        "",
    ]
    return "\n".join(lines)


def font(name: str, size: int) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/malgunbd.ttf") if name == "bold" else Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf") if name == "bold" else Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt: ImageFont.ImageFont, fill: str, width: int) -> None:
    x, y = xy
    for line in wrap(draw, text, fnt, width):
        draw.text((x, y), line, fill=fill, font=fnt)
        y += 25


def _render_legacy_revised_wbs_png_v1() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(img)
    dark = "#243044"
    mid = "#667085"
    blue = "#dcebff"
    green = "#e5f7eb"
    yellow = "#fff4d6"
    orange = "#ffe5d0"
    purple = "#eee5ff"
    header = font("bold", 38)
    sub = font("bold", 24)
    body = font("regular", 16)
    small = font("regular", 14)
    draw.text((60, 42), "[스크럼 산출물]", fill=dark, font=font("bold", 34))
    draw.text((150, 92), "WBS 및 간트차트", fill="#f00078", font=font("bold", 48))
    draw.text((610, 160), "AgendaFrame 수정 WBS 요약", fill=dark, font=body)

    headers = ["WBS", "상위 작업", "핵심 세부 작업", "담당", "산출물", "기간"]
    widths = [85, 180, 430, 95, 255, 165]
    x, y = 150, 195
    h_header = 44
    for idx, title in enumerate(headers):
        x1 = x + sum(widths[:idx])
        draw.rectangle((x1, y, x1 + widths[idx], y + h_header), fill="#d9d9d9", outline="#b8b8b8")
        draw.text((x1 + 10, y + 12), title, fill=dark, font=font("bold", 16))
    rows = [
        ["1.0", "기획·수집 기준", "MVP 범위, 언론사 3~5개, 수집 항목, 비저장 정책 확정", "공동", "수정 백로그·저장 정책", "7/7~7/12"],
        ["2.0", "기사 수집 MVP", "Playwright 수집, 3개 언론사 확장, 수집 로그·갱신 시각", "강준혁", "수집 코드·샘플 데이터", "7/10~7/31"],
        ["3.0", "저장·이슈 묶기", "메타데이터 DB, 임베딩, 유사 기사 클러스터링, 초기 검증", "공동", "DB 스키마·클러스터링 검증", "7/17~8/8"],
        ["4.0", "의제 점수·대시보드", "점수 산식, 랭킹, 이슈 상세, 원문 링크, 언론사 비교", "공동", "의제 랭킹·이슈 상세", "8/1~8/21"],
        ["5.0", "프레임·AI 리포트", "프레임 코드북, Gemini 분석, 근거 표현, AI 리포트", "공동", "프레임 비교·AI 리포트", "8/5~8/29"],
        ["6.0", "검증·발표", "MVP 통합 테스트, 오류 수정, 발표 자료, 데모 시나리오", "공동", "통합 테스트·발표 자료", "8/26~10/8"],
    ]
    cy = y + h_header
    row_h = 62
    for ridx, row in enumerate(rows):
        fill = "white" if ridx % 2 == 0 else "#fcfcfc"
        for cidx, value in enumerate(row):
            x1 = x + sum(widths[:cidx])
            draw.rectangle((x1, cy, x1 + widths[cidx], cy + row_h), fill=fill, outline="#b8b8b8")
            draw_text(draw, (x1 + 8, cy + 9), value, small, dark, widths[cidx] - 16)
        cy += row_h

    draw.text((650, 625), "수정 간트차트", fill=dark, font=body)
    gx, gy = 150, 665
    task_w = 285
    week_w = 116
    weeks = ["7/7", "7/14", "7/21", "7/28", "8/4", "8/11", "8/18", "8/25", "9/1"]
    tasks = [
        ("범위·수집 기준", 0, 1, blue),
        ("기사 수집 MVP", 1, 3, orange),
        ("저장·이슈 묶기", 1, 5, green),
        ("의제 점수·대시보드", 3, 7, yellow),
        ("프레임·AI 리포트", 4, 8, purple),
        ("검증·발표 준비", 7, 9, "#e5e7eb"),
    ]
    draw.rectangle((gx, gy, gx + task_w + week_w * len(weeks), gy + 42), fill="#d9d9d9", outline="#b8b8b8")
    draw.text((gx + 12, gy + 11), "작업", fill=dark, font=font("bold", 16))
    for idx, wk in enumerate(weeks):
        x1 = gx + task_w + idx * week_w
        draw.rectangle((x1, gy, x1 + week_w, gy + 42), fill="#d9d9d9", outline="#b8b8b8")
        draw.text((x1 + 34, gy + 12), wk, fill=dark, font=font("bold", 15))
    cy = gy + 42
    for task, start, end, color in tasks:
        draw.rectangle((gx, cy, gx + task_w, cy + 44), fill="white", outline="#dddddd")
        draw.text((gx + 12, cy + 12), task, fill=dark, font=body)
        for idx in range(len(weeks)):
            x1 = gx + task_w + idx * week_w
            draw.rectangle((x1, cy, x1 + week_w, cy + 44), fill="white", outline="#eeeeee")
        draw.rounded_rectangle(
            (gx + task_w + start * week_w + 12, cy + 10, gx + task_w + end * week_w - 12, cy + 34),
            radius=9,
            fill=color,
            outline="#8a8a8a",
        )
        cy += 44
    draw.text((60, 957), "MVP: PBL-01~17 / 후순위: PBL-18~20", fill=mid, font=small)
    draw.text((1370, 957), "AgendaFrame", fill=mid, font=small)
    img.save(OUT / "wbs_gantt.png", quality=95)


def _render_legacy_revised_wbs_png_v2() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1700, 1180), "white")
    draw = ImageDraw.Draw(img)

    navy = "#0B1F3A"
    pink = "#F00078"
    dark = "#202B3D"
    mid = "#667085"
    grid = "#D8DEE8"
    soft_blue = "#EAF2FF"
    soft_orange = "#FFF0E3"
    soft_green = "#EAF7EF"
    soft_yellow = "#FFF6DF"
    soft_purple = "#F3EDFF"
    soft_gray = "#F2F4F7"

    title_font = font("bold", 52)
    h_font = font("bold", 22)
    b_font = font("regular", 16)
    s_font = font("regular", 14)
    xs_font = font("regular", 12)

    def center(text: str, rect: tuple[int, int, int, int], fnt: ImageFont.ImageFont, fill: str = dark) -> None:
        box = draw.textbbox((0, 0), text, font=fnt)
        tw, th = box[2] - box[0], box[3] - box[1]
        draw.text((rect[0] + (rect[2] - rect[0] - tw) / 2, rect[1] + (rect[3] - rect[1] - th) / 2), text, fill=fill, font=fnt)

    def write_wrapped(x: int, y: int, text: str, fnt: ImageFont.ImageFont, fill: str, max_width: int, line_h: int = 22, max_lines: int = 2) -> None:
        lines = wrap(draw, text, fnt, max_width)
        if len(lines) > max_lines:
            lines = lines[:max_lines]
            lines[-1] = lines[-1].rstrip(" .") + "..."
        for line in lines:
            draw.text((x, y), line, fill=fill, font=fnt)
            y += line_h

    def rounded_label(rect: tuple[int, int, int, int], title: str, fill: str, outline: str) -> None:
        draw.rounded_rectangle(rect, radius=12, fill=fill, outline=outline, width=2)
        center(title, rect, s_font, dark)

    # Title area
    draw.text((65, 48), "[스크럼 산출물]", fill=dark, font=font("bold", 34))
    draw.text((150, 102), "WBS 및 간트차트", fill=pink, font=title_font)
    draw.text((154, 166), "AgendaFrame 수정 백로그 기준 일정 계획", fill=mid, font=b_font)

    meta = (1110, 60, 1608, 178)
    draw.rounded_rectangle(meta, radius=14, fill=soft_blue, outline="#C9DDF8", width=2)
    meta_items = [
        ("작성일", "2026-07-07"),
        ("작성 담당", "강준혁"),
        ("기준 백로그", "PBL-01~20"),
        ("MVP 범위", "PBL-01~17"),
    ]
    mx = 1134
    for label, value in meta_items:
        draw.text((mx, 84), label, fill=mid, font=xs_font)
        draw.text((mx, 112), value, fill=dark, font=font("bold", 15))
        mx += 118

    # WBS table
    draw.text((150, 225), "1. WBS 요약", fill=dark, font=h_font)
    draw.text((305, 229), "주요 작업 단위를 MVP 흐름에 맞춰 재정렬", fill=mid, font=s_font)

    x, y = 150, 262
    widths = [86, 180, 405, 96, 270, 135, 90]
    headers = ["WBS", "상위 작업", "핵심 세부 작업", "담당", "산출물", "기간", "구분"]
    table_w = sum(widths)
    row_h = 66
    head_h = 46
    draw.rounded_rectangle((x, y, x + table_w, y + head_h + row_h * 6), radius=10, fill="white", outline=grid, width=2)
    cx = x
    for idx, header in enumerate(headers):
        draw.rectangle((cx, y, cx + widths[idx], y + head_h), fill=navy, outline=navy)
        center(header, (cx, y, cx + widths[idx], y + head_h), font("bold", 15), "white")
        cx += widths[idx]

    rows = [
        ["1.0", "기획·수집 기준", "MVP 범위, 언론사 3~5개, 수집 항목, 비저장 정책 확정", "공동", "수정 백로그·저장 정책", "7/7~7/12", "MVP"],
        ["2.0", "기사 수집 MVP", "Playwright 수집, 3개 언론사 확장, 수집 로그·갱신 시각", "강준혁", "수집 코드·샘플 데이터", "7/10~7/31", "MVP"],
        ["3.0", "저장·이슈 묶기", "메타데이터 DB, 임베딩, 유사 기사 클러스터링, 초기 검증", "공동", "DB 스키마·클러스터링 검증", "7/17~8/8", "MVP"],
        ["4.0", "의제 점수·대시보드", "점수 산식, 랭킹, 이슈 상세, 원문 링크, 언론사 비교", "공동", "의제 랭킹·이슈 상세", "8/1~8/21", "MVP"],
        ["5.0", "프레임·AI 리포트", "프레임 코드북, Gemini 분석, 근거 표현, AI 리포트", "공동", "프레임 비교·AI 리포트", "8/5~8/29", "MVP"],
        ["6.0", "검증·발표", "MVP 통합 테스트, 오류 수정, 발표 자료, 데모 시나리오", "공동", "통합 테스트·발표 자료", "8/26~10/8", "마감"],
    ]
    fills = [soft_blue, soft_orange, soft_green, soft_yellow, soft_purple, soft_gray]
    for ridx, row in enumerate(rows):
        cy = y + head_h + ridx * row_h
        cx = x
        for cidx, value in enumerate(row):
            fill = "#FFFFFF" if ridx % 2 == 0 else "#FCFCFD"
            draw.rectangle((cx, cy, cx + widths[cidx], cy + row_h), fill=fill, outline=grid)
            if cidx == 0:
                draw.rounded_rectangle((cx + 13, cy + 20, cx + widths[cidx] - 13, cy + 46), radius=8, fill=fills[ridx], outline="#B8C0CC")
                center(value, (cx + 13, cy + 20, cx + widths[cidx] - 13, cy + 46), font("bold", 13), dark)
            elif cidx == 6:
                color = "#D92D20" if value == "마감" else "#027A48"
                draw.rounded_rectangle((cx + 12, cy + 20, cx + widths[cidx] - 12, cy + 46), radius=8, fill="#F6FEF9" if value == "MVP" else "#FFF4F4", outline=color)
                center(value, (cx + 12, cy + 20, cx + widths[cidx] - 12, cy + 46), font("bold", 13), color)
            else:
                write_wrapped(cx + 10, cy + 14, value, s_font, dark, widths[cidx] - 20, line_h=20, max_lines=2)
            cx += widths[cidx]

    # Gantt chart
    gy = 740
    draw.text((150, gy - 45), "2. 간트차트", fill=dark, font=h_font)
    draw.text((290, gy - 41), "주차별 진행 구간과 병행 작업 확인", fill=mid, font=s_font)
    gx = 150
    task_w = 285
    week_w = 105
    weeks = ["7/7", "7/14", "7/21", "7/28", "8/4", "8/11", "8/18", "8/25", "9/1"]
    chart_w = task_w + week_w * len(weeks)
    draw.rounded_rectangle((gx, gy, gx + chart_w, gy + 44 + 56 * 6), radius=10, fill="white", outline=grid, width=2)
    draw.rectangle((gx, gy, gx + chart_w, gy + 44), fill=navy)
    draw.text((gx + 18, gy + 13), "작업", fill="white", font=font("bold", 15))
    for idx, wk in enumerate(weeks):
        x1 = gx + task_w + idx * week_w
        draw.line((x1, gy, x1, gy + 44 + 56 * 6), fill=grid, width=1)
        center(wk, (x1, gy, x1 + week_w, gy + 44), font("bold", 15), "white")
    draw.line((gx + chart_w, gy, gx + chart_w, gy + 44 + 56 * 6), fill=grid, width=1)

    tasks = [
        ("범위·수집 기준", "PBL-01~05", 0, 1, soft_blue, "#2B6BD2"),
        ("기사 수집 MVP", "PBL-03~04", 1, 3, soft_orange, "#E87511"),
        ("저장·이슈 묶기", "PBL-05~07", 1, 5, soft_green, "#168356"),
        ("의제 점수·대시보드", "PBL-08~12", 3, 7, soft_yellow, "#D8A300"),
        ("프레임·AI 리포트", "PBL-13~17", 4, 8, soft_purple, "#7A4DD8"),
        ("검증·발표 준비", "통합 테스트", 7, 9, soft_gray, "#667085"),
    ]
    for idx, (task, pbl, start, end, fill, outline) in enumerate(tasks):
        cy = gy + 44 + idx * 56
        draw.rectangle((gx, cy, gx + chart_w, cy + 56), fill="#FFFFFF" if idx % 2 == 0 else "#FCFCFD", outline=grid)
        draw.text((gx + 18, cy + 11), task, fill=dark, font=font("bold", 15))
        draw.text((gx + 18, cy + 33), pbl, fill=mid, font=xs_font)
        bar_x1 = gx + task_w + start * week_w + 16
        bar_x2 = gx + task_w + end * week_w - 16
        draw.rounded_rectangle((bar_x1, cy + 16, bar_x2, cy + 40), radius=10, fill=fill, outline=outline, width=2)
        center(f"{weeks[start]}~{weeks[end - 1]}", (bar_x1, cy + 16, bar_x2, cy + 40), xs_font, dark)

    # Legend and notes
    legend_x = 1415
    legend_y = 742
    draw.rounded_rectangle((legend_x, legend_y, 1608, 918), radius=12, fill="#F8FAFC", outline=grid, width=2)
    draw.text((legend_x + 22, legend_y + 18), "범례 / 기준", fill=dark, font=font("bold", 18))
    notes = [
        ("MVP", "PBL-01~17 구현"),
        ("후순위", "PBL-18~20 별도 스프린트"),
        ("저장 정책", "기사 전문 미저장"),
        ("검증 목표", "클러스터링 정합성 80%+"),
    ]
    ny = legend_y + 55
    for label, value in notes:
        draw.text((legend_x + 24, ny), label, fill=mid, font=xs_font)
        draw.text((legend_x + 104, ny), value, fill=dark, font=xs_font)
        ny += 28

    draw.line((60, 1135, 1640, 1135), fill="#EEF2F7", width=2)
    draw.text((60, 1150), "※ 수정 백로그 기준: PBL-01~17은 MVP, PBL-18 상세 검색·PBL-19 결과 공유/PDF·PBL-20 사용자 편의는 후순위 확장", fill=mid, font=xs_font)
    draw.text((1515, 1150), "AgendaFrame", fill=mid, font=xs_font)
    img.save(OUT / "wbs_gantt.png", quality=95)


def render_wbs_summary_tree() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), "#F6F8FA")
    draw = ImageDraw.Draw(img)

    navy = "#0B1F3A"
    dark = "#202B3D"
    mid = "#667085"
    line = "#9AA7B8"
    card_shadow = "#E2E8F0"
    colors = ["#2B6BD2", "#0E9F6E", "#F97316", "#DC2626", "#475569", "#0F9996", "#7C3AED"]
    fills = ["#EFF6FF", "#ECFDF5", "#FFF7ED", "#FEF2F2", "#F8FAFC", "#ECFEFF", "#F5F3FF"]

    title = font("bold", 34)
    section = font("bold", 20)
    body = font("regular", 15)
    small = font("regular", 13)

    def center(text: str, rect: tuple[int, int, int, int], fnt: ImageFont.ImageFont, fill: str = dark) -> None:
        box = draw.textbbox((0, 0), text, font=fnt)
        tw, th = box[2] - box[0], box[3] - box[1]
        draw.text((rect[0] + (rect[2] - rect[0] - tw) / 2, rect[1] + (rect[3] - rect[1] - th) / 2), text, fill=fill, font=fnt)

    def connector(points: list[tuple[int, int]]) -> None:
        for p1, p2 in zip(points, points[1:]):
            draw.line((p1, p2), fill=line, width=3)

    def card(rect: tuple[int, int, int, int], idx: str, title_text: str, bullets: list[str], accent: str, fill: str) -> None:
        sx1, sy1, sx2, sy2 = rect[0] + 8, rect[1] + 10, rect[2] + 8, rect[3] + 10
        draw.rounded_rectangle((sx1, sy1, sx2, sy2), radius=16, fill=card_shadow)
        draw.rounded_rectangle(rect, radius=16, fill="white", outline="#E5EAF2", width=1)
        draw.rounded_rectangle((rect[0], rect[1], rect[0] + 8, rect[3]), radius=5, fill=accent)
        badge = (rect[0] + 24, rect[1] + 24, rect[0] + 76, rect[1] + 56)
        draw.rounded_rectangle(badge, radius=16, fill=fill, outline=accent, width=2)
        center(idx, badge, font("bold", 14), accent)
        draw.text((rect[0] + 92, rect[1] + 26), title_text, fill=dark, font=section)
        by = rect[1] + 70
        for bullet in bullets:
            draw.ellipse((rect[0] + 34, by + 6, rect[0] + 42, by + 14), fill=accent)
            draw.text((rect[0] + 58, by), bullet, fill=mid, font=body)
            by += 27

    # Header
    draw.rounded_rectangle((70, 45, 1530, 150), radius=20, fill="white", outline="#E5EAF2", width=1)
    draw.text((105, 68), "AgendaFrame WBS 요약 트리", fill=dark, font=title)
    draw.text((106, 112), "수정 백로그를 7개 Work Package 중심으로 구조화", fill=mid, font=body)
    draw.rounded_rectangle((1285, 72, 1490, 112), radius=22, fill="#EEF2FF", outline="#CDD5FF", width=1)
    center("WBS Summary", (1285, 72, 1490, 112), font("bold", 14), "#3730A3")

    # Root
    root = (540, 190, 1060, 285)
    draw.rounded_rectangle((root[0] + 8, root[1] + 10, root[2] + 8, root[3] + 10), radius=18, fill=card_shadow)
    draw.rounded_rectangle(root, radius=18, fill=navy)
    center("AgendaFrame", (root[0], root[1] + 18, root[2], root[1] + 52), font("bold", 28), "white")
    center("AI 기반 뉴스 의제·프레임 분석 플랫폼", (root[0], root[1] + 55, root[2], root[1] + 86), body, "#DCE6F3")

    top_cards = [
        ((90, 365, 405, 535), "01", "기획/요구사항", ["MVP 범위 확정", "언론사 3~5개 선정", "기사 전문 비저장 정책"], colors[0], fills[0]),
        ((455, 365, 770, 535), "02", "기사 수집 MVP", ["홈페이지 구조 조사", "Playwright 수집 구현", "수집 로그·갱신 시각"], colors[1], fills[1]),
        ((820, 365, 1135, 535), "03", "데이터/이슈", ["메타데이터 DB", "임베딩·클러스터링", "초기 정합성 검증"], colors[2], fills[2]),
        ((1185, 365, 1500, 535), "04", "의제/대시보드", ["의제 점수 산식", "랭킹·이슈 상세", "언론사별 보도 비교"], colors[3], fills[3]),
    ]
    bottom_cards = [
        ((240, 620, 555, 790), "05", "AI/프레임", ["프레임 코드북", "Gemini 분석", "근거 표현·AI 리포트"], colors[4], fills[4]),
        ((645, 620, 960, 790), "06", "검증/시연", ["클러스터링 검증", "프레임 수동 검토", "MVP 통합 테스트"], colors[5], fills[5]),
        ((1050, 620, 1365, 790), "07", "문서/발표", ["통합 산출물 보고서", "발표 자료", "데모 시나리오"], colors[6], fills[6]),
    ]

    # Connectors
    root_mid = ((root[0] + root[2]) // 2, root[3])
    top_bus_y = 325
    bottom_bus_y = 585
    connector([root_mid, (root_mid[0], top_bus_y), (255, top_bus_y), (1338, top_bus_y)])
    for rect, *_ in top_cards:
        cx = (rect[0] + rect[2]) // 2
        connector([(cx, top_bus_y), (cx, rect[1])])
    connector([(800, 535), (800, bottom_bus_y), (397, bottom_bus_y), (1208, bottom_bus_y)])
    for rect, *_ in bottom_cards:
        cx = (rect[0] + rect[2]) // 2
        connector([(cx, bottom_bus_y), (cx, rect[1])])

    for args in top_cards + bottom_cards:
        card(*args)

    draw.text((70, 850), "※ PBL-01~17은 MVP, PBL-18~20은 후순위 확장 기능으로 관리", fill=mid, font=small)
    draw.text((1442, 850), "AgendaFrame", fill=mid, font=small)
    img.save(OUT / "wbs_summary_tree.png", quality=95)


def render_revised_wbs_png(output_dir: Path | None = None) -> None:
    global OUT

    if output_dir is not None:
        OUT = output_dir.resolve()
    OUT.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1700, 1080), "#F6F8FA")
    draw = ImageDraw.Draw(img)

    navy = "#0B1F3A"
    dark = "#202B3D"
    mid = "#667085"
    grid = "#E6EBF2"
    blue = "#2B6BD2"
    green = "#059669"
    orange = "#EA580C"
    red = "#DC2626"
    purple = "#7C3AED"
    slate = "#475569"
    teal = "#0F9996"
    pink = "#BE185D"
    group_fills = {
        "기획/요구사항": "#EAF2FF",
        "수집": "#ECFDF5",
        "데이터/이슈": "#FFF7ED",
        "의제/대시보드": "#FEF2F2",
        "AI/프레임": "#F5F3FF",
        "검증/시연": "#ECFEFF",
        "문서/발표": "#FDF2F8",
    }

    title_font = font("bold", 36)
    h_font = font("bold", 16)
    b_font = font("regular", 14)
    small = font("regular", 12)
    tiny = font("regular", 11)

    def center(text: str, rect: tuple[int, int, int, int], fnt: ImageFont.ImageFont, fill: str = dark) -> None:
        box = draw.textbbox((0, 0), text, font=fnt)
        tw, th = box[2] - box[0], box[3] - box[1]
        draw.text((rect[0] + (rect[2] - rect[0] - tw) / 2, rect[1] + (rect[3] - rect[1] - th) / 2), text, fill=fill, font=fnt)

    def date_to_x(label: str) -> int:
        month, day = map(int, label.split("/"))
        days_by_month = {7: 0, 8: 31, 9: 62, 10: 92}
        day_index = days_by_month[month] + day - 7
        return int(timeline_x1 + day_index / total_days * timeline_w)

    def bar(y: int, start: str, end: str, color: str) -> None:
        x1 = date_to_x(start)
        x2 = date_to_x(end)
        x2 = max(x2, x1 + 70)
        draw.rounded_rectangle((x1, y + 10, x2, y + 31), radius=10, fill=color)
        center(f"{start}-{end}", (x1, y + 10, x2, y + 31), tiny, "white")

    # Header
    draw.text((90, 52), "AgendaFrame WBS 기반 간트차트", fill=dark, font=title_font)
    draw.text((92, 96), "기간: 7월 1주차부터 10월 2주차까지", fill=mid, font=b_font)

    card = (60, 135, 1640, 1000)
    draw.rounded_rectangle((card[0] + 8, card[1] + 10, card[2] + 8, card[3] + 10), radius=22, fill="#E2E8F0")
    draw.rounded_rectangle(card, radius=22, fill="white", outline="#E5EAF2", width=1)

    table_x = 85
    table_y = 170
    group_w = 165
    package_w = 360
    timeline_x1 = table_x + group_w + package_w
    timeline_x2 = 1590
    timeline_w = timeline_x2 - timeline_x1
    total_days = 94
    header_h = 42
    row_h = 43

    # Month bands
    draw.rounded_rectangle((table_x, table_y, table_x + group_w, table_y + header_h), radius=14, fill="#F1F5F9")
    draw.rounded_rectangle((table_x + group_w, table_y, table_x + group_w + package_w, table_y + header_h), radius=14, fill="#F1F5F9")
    center("구분", (table_x, table_y, table_x + group_w, table_y + header_h), h_font, "#344054")
    center("Work Package", (table_x + group_w, table_y, timeline_x1, table_y + header_h), h_font, "#344054")

    months = [
        ("7월", "7/7", "7/31", "#EEF4FF"),
        ("8월", "8/1", "8/31", "#ECFDF3"),
        ("9월", "9/1", "9/30", "#FFF7E6"),
        ("10월", "10/1", "10/8", "#FDF2F8"),
    ]
    for name, start, end, fill in months:
        x1 = date_to_x(start)
        x2 = date_to_x(end)
        draw.rounded_rectangle((x1 + 4, table_y + 5, x2 - 4, table_y + header_h - 5), radius=16, fill=fill, outline="#D8E0EC")
        center(name, (x1, table_y + 5, x2, table_y + header_h - 5), h_font, "#344054")

    weeks = ["7/7", "7/14", "7/21", "7/28", "8/4", "8/11", "8/18", "8/25", "9/1", "9/8", "9/15", "9/22", "9/29", "10/6"]
    week_y = table_y + header_h
    draw.line((timeline_x1, week_y, timeline_x2, week_y), fill=grid, width=2)
    for wk in weeks:
        x = date_to_x(wk)
        draw.line((x, week_y, x, card[3] - 45), fill=grid, width=1)
        center(wk, (x - 25, week_y + 4, x + 25, week_y + 24), tiny, mid)

    rows = [
        ("기획/요구사항", "1.1 MVP 범위 및 백로그 확정", "7/7", "7/8", blue),
        ("", "1.2 수집 대상/정책 분야 정의", "7/7", "7/11", blue),
        ("", "1.3 기사 전문 비저장 정책 정리", "7/9", "7/12", blue),
        ("수집", "2.1 언론사 홈페이지 구조 조사", "7/10", "7/14", green),
        ("", "2.2 Playwright 수집 MVP 구현", "7/14", "7/23", green),
        ("", "2.3 3개 언론사 확장 및 로그 구현", "7/21", "7/31", green),
        ("데이터/이슈", "3.1 메타데이터 DB 및 저장 로직", "7/17", "7/26", orange),
        ("", "3.2 임베딩·유사 기사 클러스터링", "7/24", "8/5", orange),
        ("", "3.3 클러스터링 초기 검증", "8/1", "8/8", orange),
        ("의제/대시보드", "4.1 의제 점수 산식 및 랭킹 구현", "8/1", "8/14", red),
        ("", "4.2 이슈 상세·원문 링크·언론사 비교", "8/10", "8/21", red),
        ("AI/프레임", "5.1 프레임 코드북·Gemini 분석", "8/5", "8/22", purple),
        ("", "5.2 프레임 비교·근거 표시·수동 검증", "8/18", "8/28", purple),
        ("", "5.3 AI 리포트 구현", "8/20", "8/29", purple),
        ("검증/시연", "6.1 MVP 통합 테스트", "8/26", "8/31", teal),
        ("문서/발표", "7.1 발표 자료 및 데모 시나리오", "9/1", "10/8", pink),
    ]
    first_row_y = week_y + 26
    for idx, (group, package, start, end, color) in enumerate(rows):
        y = first_row_y + idx * row_h
        fill = "#FBFCFE" if idx % 2 == 0 else "#FFFFFF"
        draw.rectangle((table_x, y, timeline_x2, y + row_h), fill=fill)
        draw.line((table_x, y + row_h, timeline_x2, y + row_h), fill="#F0F3F7", width=1)
        if group:
            pill = (table_x + 18, y + 9, table_x + group_w - 18, y + 31)
            draw.rounded_rectangle(pill, radius=12, fill=group_fills[group], outline="#DDE5EF")
            center(group, pill, tiny, dark)
        draw.text((table_x + group_w + 20, y + 12), package, fill=dark, font=b_font)
        bar(y, start, end, color)

    draw.text((85, 958), "MVP: PBL-01~17 / 후순위: PBL-18~20", fill=mid, font=small)
    draw.text((1450, 958), "AgendaFrame", fill=mid, font=small)
    img.save(OUT / "wbs_gantt.png", quality=95)
    render_wbs_summary_tree()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, size: int = 10) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=size)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=10)


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if parts and all(re.fullmatch(r":?-{3,}:?", part.replace(" ", "")) for part in parts):
            continue
        rows.append(parts)
    return rows


def add_markdown_table_as_entries(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = rows[0]
    for idx, row in enumerate(rows[1:], 1):
        label = row[0] if row and row[0] else f"항목 {idx}"
        p = doc.add_paragraph()
        r = p.add_run(str(label))
        set_run_font(r, size=9, bold=True, color="1F4D78")
        details = []
        for h, v in zip(headers[1:], row[1:]):
            if str(v).strip():
                details.append(f"{h}: {v}")
        r2 = p.add_run(" - " + " / ".join(details))
        set_run_font(r2, size=8)


def add_small_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=8, bold=(r_idx == 0), color="FFFFFF" if r_idx == 0 else None)
            if r_idx == 0:
                set_cell_shading(cell, "0B1F3A")
    doc.add_paragraph()


def add_markdown(doc: Document, path: Path, title: str, max_lines: int | None = None) -> None:
    doc.add_heading(title, level=1)
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    if max_lines is not None:
        lines = lines[:max_lines]
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            add_para(doc, "\n".join(block), size=8)
            i += 1
            continue
        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_markdown_table(table_lines)
            if rows and len(rows[0]) <= 5 and len(rows) <= 18:
                add_small_table(doc, rows)
            else:
                add_markdown_table_as_entries(doc, rows)
            continue
        if line.startswith("#"):
            level = min(line.count("#", 0, len(line) - len(line.lstrip("#"))) + 1, 3)
            heading_text = line.lstrip("#").strip()
            if heading_text:
                doc.add_heading(heading_text, level=level)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            set_run_font(run, size=10)
            i += 1
            continue
        add_para(doc, line.strip(), size=10)
        i += 1


def add_image(doc: Document, path: Path, title: str) -> None:
    if not path.exists():
        return
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.5))


def add_xlsx_summary(doc: Document, path: Path) -> None:
    from openpyxl import load_workbook

    doc.add_heading("기능 명세서", level=1)
    add_para(doc, "기능 명세서는 별도 XLSX 파일로 관리하며, 통합 보고서에는 주요 검토 관점과 파일 구성만 요약한다.", size=10)
    if not path.exists():
        return
    wb = load_workbook(path, data_only=True)
    rows = [["시트명", "행 수", "열 수", "비고"]]
    for ws in wb.worksheets:
        rows.append([ws.title, str(ws.max_row), str(ws.max_column), "원본 XLSX 별도 첨부"])
    add_small_table(doc, rows)


def build_docx(output_path: Path = DEFAULT_DOCX_OUT) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("2E74B5")
    styles["Heading 2"].font.color.rgb = RGBColor.from_string("1F4D78")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AgendaFrame 통합 산출물 보고서")
    set_run_font(r, size=24, bold=True, color="0B1F3A")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("선행연구 · 기술스택/아키텍처 · 기능명세 · 백로그 · WBS/간트차트 · UML 통합본")
    set_run_font(r, size=11, color="475467")
    add_para(doc, "작성일: 2026-07-07 / 프로젝트명: AgendaFrame / 작성 담당: 강준혁", size=10)

    doc.add_heading("1. 검토용 요약", level=1)
    add_bullets(
        doc,
        [
            "AgendaFrame은 주요 언론사의 홈페이지 배치와 보도 빈도를 기반으로 오늘의 공적 의제를 산출하고, 언론사별 관점/프레임 차이를 비교하는 AI 기반 뉴스 의제·프레임 분석 플랫폼이다.",
            "수정 백로그 기준 MVP 범위는 PBL-01~PBL-17이며, 상세 검색·PDF 공유·즐겨찾기/다크모드는 후순위 확장으로 분리했다.",
            "저작권 위험을 줄이기 위해 기사 전문은 저장·재게시하지 않고 제목, URL, 언론사명, 섹션, 배치 위치, 수집 시각, 분석 결과 중심으로 저장한다.",
            "기술 구조는 Presentation, Application/API, Collection Batch, Analysis/AI, Data Layer로 구분하고 처리 파이프라인은 Collect → Extract → Store → Cluster → Score → Frame → Report → Dashboard로 정리했다.",
        ],
    )

    doc.add_heading("2. 산출물 구성", level=1)
    inventory = [
        ["구분", "파일", "설명"],
        ["선행연구", "docs/research/prior-research-and-services.md", "관련 연구와 선행 서비스를 비교해 문제 정의와 차별점을 정리"],
        ["기술 구조", "outputs/system_architecture.png", "레이어 기준 시스템 아키텍처와 처리 파이프라인 시각화"],
        ["기능 명세", "docs/specs/feature-spec.xlsx", "기능별 요구사항과 화면/처리 기준을 표 형태로 관리"],
        ["백로그", "docs/planning/product-backlog.md / docs/planning/sprint-backlog.md", "PBL-01~20과 스프린트별 구현 작업 정리"],
        ["WBS", "docs/planning/wbs-gantt.md / outputs/wbs_gantt.png", "일정, 담당, 산출물, 선행 작업, 간트차트 정리"],
        ["UML", "docs/specs/uml.md / outputs/*.png", "유스케이스, 액티비티, 클래스, 시퀀스 다이어그램 정리"],
    ]
    add_small_table(doc, inventory)

    doc.add_heading("3. 기술스택 및 시스템 아키텍처 설명", level=1)
    add_bullets(
        doc,
        [
            "프론트엔드는 React와 Firebase Hosting을 기준으로 대시보드 화면을 제공한다.",
            "API는 Cloud Run API가 담당하며 의제 목록, 이슈 상세, 언론사 비교, AI 리포트 조회를 오케스트레이션한다.",
            "기사 수집은 Cloud Scheduler와 Cloud Run Jobs, Playwright 조합으로 정기 실행한다.",
            "BigQuery는 기사 메타데이터, 이슈, 의제 점수, 프레임 분석 결과를 저장하며 BigQuery Vector Search는 유사 기사 묶기에 활용한다.",
            "Vertex AI Gemini는 프레임 분석과 AI 리포트 생성을 담당한다.",
        ],
    )
    add_image(doc, OUT / "system_architecture.png", "시스템 아키텍처")
    add_image(doc, OUT / "wbs_summary_tree.png", "WBS 요약 트리")
    add_image(doc, OUT / "wbs_gantt.png", "WBS 및 간트차트")

    add_xlsx_summary(doc, SPECS / "feature-spec.xlsx")

    doc.add_heading("4. UML 시각 자료", level=1)
    add_image(doc, OUT / "usecase_diagram.png", "유스케이스 다이어그램")
    add_image(doc, OUT / "activity_diagram.png", "액티비티 다이어그램")
    add_image(doc, OUT / "class_diagram_domain.png", "클래스 다이어그램 1: 도메인 모델 및 데이터 관계")
    add_image(doc, OUT / "class_diagram_implementation.png", "클래스 다이어그램 2: 애플리케이션 구현 구조")
    add_image(doc, OUT / "sequence_uc01_uc02.png", "시퀀스 다이어그램: UC-01~UC-02 의제 조회 및 이슈 상세 조회")
    add_image(doc, OUT / "sequence_uc03_uc05.png", "시퀀스 다이어그램: UC-03~UC-05 비교 분석 및 AI 리포트")
    add_image(doc, OUT / "sequence_uc06.png", "시퀀스 다이어그램: UC-06 PDF 내보내기")
    add_image(doc, OUT / "sequence_uc07.png", "시퀀스 다이어그램: UC-07 기사 자동 수집 및 분석")

    doc.add_page_break()
    doc.add_heading("부록: 원문 문서", level=1)
    for path, title in [
        (PLANNING / "product-backlog.md", "부록 A. 프로덕트 백로그"),
        (PLANNING / "sprint-backlog.md", "부록 B. 스프린트 백로그"),
        (PLANNING / "wbs-gantt.md", "부록 C. WBS 및 간트차트"),
        (RESEARCH / "prior-research-and-services.md", "부록 D. 선행연구 및 선행서비스 검토"),
        (SPECS / "usecase-spec.md", "부록 E. 유스케이스 명세서"),
        (SPECS / "uml.md", "부록 F. UML 문서"),
    ]:
        if path.exists():
            add_markdown(doc, path, title)
            doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build AgendaFrame submission artifacts.")
    parser.add_argument(
        "--asset-dir",
        type=Path,
        default=DEFAULT_OUT,
        help="Directory containing and receiving submission images.",
    )
    parser.add_argument(
        "--docx-output",
        type=Path,
        default=DEFAULT_DOCX_OUT,
        help="Path for the generated integrated DOCX report.",
    )
    parser.add_argument(
        "--backlog-source",
        type=Path,
        help="Optional UTF-8 source used to regenerate backlog Markdown files.",
    )
    return parser.parse_args()


def main() -> None:
    global OUT

    args = parse_args()
    OUT = args.asset_dir.resolve()
    PLANNING.mkdir(parents=True, exist_ok=True)
    SPECS.mkdir(parents=True, exist_ok=True)
    RESEARCH.mkdir(parents=True, exist_ok=True)
    SUBMISSION.mkdir(parents=True, exist_ok=True)
    OUT.mkdir(parents=True, exist_ok=True)
    if args.backlog_source:
        backlog_source = args.backlog_source.resolve()
        if not backlog_source.is_file():
            raise FileNotFoundError(f"backlog source not found: {backlog_source}")
        text = backlog_source.read_text(encoding="utf-8")
        product, sprint, _ = split_revision_backlog(text)
        (PLANNING / "product-backlog.md").write_text(product, encoding="utf-8")
        (PLANNING / "sprint-backlog.md").write_text(sprint, encoding="utf-8")
        (PLANNING / "wbs-gantt.md").write_text(build_wbs_markdown(), encoding="utf-8")
    render_revised_wbs_png()
    docx_output = args.docx_output.resolve()
    build_docx(docx_output)
    print(f"wrote submission assets to {OUT} and report to {docx_output}")


if __name__ == "__main__":
    main()
